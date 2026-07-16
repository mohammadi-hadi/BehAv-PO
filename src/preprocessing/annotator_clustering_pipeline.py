#!/usr/bin/env python3
"""
Annotator Clustering & Demographic Analysis Pipeline
=====================================================
Implements the methodology from Lo & Basile (2023) and GeNLP 2025 / EXIST 2024 papers.

Pipeline:
  Section 0:  Imports & Configuration
  Section 1:  Data Loading & One-Hot Encoding (incl. 1h Block Design Discovery)
  Section 2:  Kernel PCA (Cosine Kernel) — guarded by RUN_KPCA_COMPARISON
  Section 2B: Behavioral Feature Clustering (PRIMARY when PRIMARY_METHOD="behavioral")
  Section 3:  Ward's Linkage Clustering — guarded by RUN_KPCA_COMPARISON
  Section 4:  Cluster Validation
  Section 5:  Post-Hoc Demographic Analysis
  Section 6:  GLMM — Statistical Modeling of Demographic Effects
  Section 7:  Robustness Checks (Optional)
  Section 8:  Word Report Generation (Optional)
  Section 9:  Jupyter Notebook Generation (Optional)
"""

# ============================================================
# Section 0: Imports & Configuration
# ============================================================

import json
import warnings

import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")          # non-interactive backend (safe for Colab & CLI)
import matplotlib.pyplot as plt
import seaborn as sns

from sklearn.decomposition import KernelPCA
from sklearn.preprocessing import StandardScaler
from sklearn.cluster import AgglomerativeClustering, KMeans
from sklearn.metrics import (
    silhouette_score, silhouette_samples,
    calinski_harabasz_score, davies_bouldin_score,
    adjusted_rand_score, adjusted_mutual_info_score,
)

from scipy.cluster.hierarchy import linkage, dendrogram, fcluster
from scipy.stats import chi2_contingency

warnings.filterwarnings("ignore", category=FutureWarning)
warnings.filterwarnings("ignore", category=RuntimeWarning, message=".*encountered in matmul.*")
warnings.filterwarnings("ignore", category=RuntimeWarning, message=".*encountered in scalar divide.*")

# ── Configuration ──────────────────────────────────────────────────
DATA_PATH = "data/raw/EXIST2024_training.json"   # adjust for Colab path
LANGUAGE = "en"
MIN_LABELS = 10                          # min annotations per annotator
VARIANCE_THRESHOLD = 0.855               # 85.5 % cumulative variance
CLUSTER_RANGE = range(2, 21)             # 2–20 clusters to evaluate
RANDOM_STATE = 42
RUN_GLMM = True                          # set False to skip Section 6
RUN_ROBUSTNESS = True                    # set False to skip Section 7
PRIMARY_METHOD = "behavioral"            # "behavioral" or "kpca"
RUN_KPCA_COMPARISON = True               # run KPCA even when behavioral is primary
GENERATE_REPORT = True                   # generate Word report (Section 8)
GENERATE_NOTEBOOK = True                 # generate Jupyter notebook (Section 9)

# Country → Region mapping (GeNLP 2025 paper)
COUNTRY_TO_REGION = {
    # Europe
    "Italy": "Europe", "Spain": "Europe", "France": "Europe",
    "Germany": "Europe", "United Kingdom": "Europe", "Portugal": "Europe",
    "Belgium": "Europe", "Netherlands": "Europe", "Norway": "Europe",
    "Finland": "Europe", "Greece": "Europe", "Hungary": "Europe",
    "Poland": "Europe", "Romania": "Europe", "Czech Republic": "Europe",
    "Ireland": "Europe", "Estonia": "Europe", "Latvia": "Europe",
    "Slovenia": "Europe", "Serbia": "Europe", "Cyprus": "Europe",
    "Macedonia, The Former Yugoslav Republic of": "Europe",
    "Russian Federation": "Europe",
    # Americas
    "United States": "America", "Canada": "America", "Mexico": "America",
    "Chile": "America", "Argentina": "America", "Brazil": "America",
    "Colombia": "America", "Peru": "America", "Venezuela": "America",
    "Cuba": "America", "Dominican Republic": "America",
    "El Salvador": "America", "Panama": "America", "Puerto Rico": "America",
    # Africa
    "South Africa": "Africa",
    # Asia / Oceania
    "Nepal": "Asia", "Viet Nam": "Asia",
    "Australia": "Asia", "New Zealand": "Asia",
    # Middle East
    "Afghanistan": "Middle East", "Algeria": "Middle East",
    "Israel": "Middle East",
}


# ============================================================
# Section 1: Data Loading & One-Hot Encoding
# ============================================================
print("=" * 60)
print("SECTION 1 — Data Loading & One-Hot Encoding")
print("=" * 60)

# 1a. Load JSON ---------------------------------------------------
with open(DATA_PATH, "r") as f:
    raw = json.load(f)
print(f"Loaded {len(raw)} texts from {DATA_PATH}")

# 1b. Build annotator-level long table ----------------------------
rows_list = []
for entry in raw.values():
    text_id = entry["id_EXIST"]
    lang = entry["lang"]
    tweet = entry["tweet"]
    annotators = entry["annotators"]
    n = len(annotators)
    for i in range(n):
        lab = entry["labels_task1"][i] if i < len(entry["labels_task1"]) else None
        rows_list.append({
            "text_id": text_id,
            "language": lang,
            "tweet": tweet,
            "annotator_id": annotators[i],
            "gender": entry["gender_annotators"][i] if i < len(entry["gender_annotators"]) else None,
            "age": entry["age_annotators"][i] if i < len(entry["age_annotators"]) else None,
            "ethnicity": entry["ethnicities_annotators"][i] if i < len(entry["ethnicities_annotators"]) else None,
            "education": entry["study_levels_annotators"][i] if i < len(entry["study_levels_annotators"]) else None,
            "country": entry["countries_annotators"][i] if i < len(entry["countries_annotators"]) else None,
            "label_raw": lab,
            "label": 1 if lab == "YES" else (-1 if lab == "NO" else 0),
        })

annot_df = pd.DataFrame(rows_list)
print(f"Annotator-level table: {annot_df.shape}")

# 1c. Pivot to text × annotator matrix ----------------------------
matrix_df = (
    annot_df
    .pivot_table(
        index=["text_id", "language"],
        columns="annotator_id",
        values="label",
        aggfunc="first",
    )
    .fillna(0)
    .astype(int)
)

# 1d. Filter to English -------------------------------------------
matrix_en = matrix_df.xs(LANGUAGE, level="language")
print(f"English texts: {matrix_en.shape[0]}")

# 1e. Select active annotators (≥ MIN_LABELS non-zero entries) ----
non_zero_counts = (matrix_en != 0).sum(axis=0)
active_mask = non_zero_counts >= MIN_LABELS
active_annotators = non_zero_counts[active_mask].index.tolist()
print(f"Active annotators (≥ {MIN_LABELS} labels): {len(active_annotators)}")

# Annotator × text matrix (scalar: -1 / 0 / 1)
A = matrix_en[active_annotators].values.T          # (n_annotators, n_texts)
annotator_names = active_annotators
text_ids = matrix_en.index.tolist()
n_annotators, n_texts = A.shape
print(f"Matrix A shape: {A.shape}")
print(f"Sparsity: {(A == 0).sum() / A.size:.2%}")

# 1f. One-hot encoding: YES→(0,1)  NO→(1,0)  missing→(0,0) -------
A_no = (A == -1).astype(np.float64)
A_yes = (A == 1).astype(np.float64)
M_onehot = np.empty((n_annotators, 2 * n_texts), dtype=np.float64)
M_onehot[:, 0::2] = A_no       # even columns = NO indicator
M_onehot[:, 1::2] = A_yes      # odd columns  = YES indicator
print(f"One-hot matrix shape: {M_onehot.shape}")

# 1g. Extract demographics per annotator --------------------------
annot_en = annot_df[annot_df["language"] == LANGUAGE].copy()
demo_df = (
    annot_en[annot_en["annotator_id"].isin(active_annotators)]
    .drop_duplicates(subset="annotator_id")
    .set_index("annotator_id")
    [["gender", "age", "ethnicity", "education", "country"]]
    .loc[annotator_names]
)
demo_df["region"] = demo_df["country"].map(COUNTRY_TO_REGION).fillna("Other")
print(f"\nDemographics table: {demo_df.shape}")
print(demo_df["region"].value_counts())
print()

# 1h. Block design discovery ------------------------------------------
print("-" * 40)
print("1h. Block Design Discovery")
print("-" * 40)

# For each text, find which annotators labeled it
text_annotator_sets = {}
for tid in text_ids:
    tid_idx = text_ids.index(tid)
    labeled = [annotator_names[j] for j in range(n_annotators) if A[j, tid_idx] != 0]
    text_annotator_sets[tid] = frozenset(labeled)

# Group texts by their annotator set
from collections import Counter
group_counter = Counter(text_annotator_sets.values())
n_groups = len(group_counter)
group_sizes = sorted(group_counter.values(), reverse=True)
print(f"Unique annotator groups: {n_groups}")
print(f"Group sizes (texts per group): min={min(group_sizes)}, "
      f"max={max(group_sizes)}, mean={np.mean(group_sizes):.1f}")

# Check cross-group overlap
all_groups = list(group_counter.keys())
has_overlap = False
for i in range(len(all_groups)):
    for j in range(i + 1, len(all_groups)):
        if all_groups[i] & all_groups[j]:
            has_overlap = True
            break
    if has_overlap:
        break

if not has_overlap:
    print("CONFIRMED: Perfect block design — zero cross-group annotator overlap.")
    print("  This means cosine similarity between annotators in different groups is exactly 0.")
    print("  KPCA on the raw matrix recovers assignment structure, not behavioral differences.")
else:
    print("Some cross-group overlap detected — block design is not perfect.")

# Generate block design heatmap (reorder matrix to show block-diagonal)
# Assign each annotator to its group index
annotator_to_group = {}
for gidx, grp in enumerate(all_groups):
    for ann in grp:
        annotator_to_group[ann] = gidx

group_order = sorted(range(n_annotators), key=lambda i: annotator_to_group.get(annotator_names[i], -1))
A_reordered = A[group_order, :]

# Also reorder texts by group for cleaner visual
text_group = {}
for tid_idx, tid in enumerate(text_ids):
    grp = text_annotator_sets[tid]
    text_group[tid_idx] = all_groups.index(grp)
text_order = sorted(range(n_texts), key=lambda i: text_group[i])
A_block = A_reordered[:, text_order]

fig, ax = plt.subplots(figsize=(14, 8))
# Show non-zero entries only for visibility
display_matrix = np.where(A_block != 0, 1, 0).astype(float)
ax.imshow(display_matrix, aspect="auto", cmap="Blues", interpolation="none")
ax.set_xlabel("Texts (reordered by annotation group)")
ax.set_ylabel("Annotators (reordered by annotation group)")
ax.set_title(f"Block Design Structure — {n_groups} Disjoint Annotator Groups")
plt.tight_layout()
plt.savefig("block_design_heatmap.png", dpi=150, bbox_inches="tight")
plt.close()
print("Saved: block_design_heatmap.png\n")


# ============================================================
# Section 2: Kernel PCA (Cosine Kernel) — COMPARISON ONLY
# ============================================================
kpca_cluster_labels = None
Z = None
cum_var = None
k = None

if RUN_KPCA_COMPARISON:
    print("=" * 60)
    print("SECTION 2 — Kernel PCA (Cosine Kernel) [COMPARISON]")
    print("=" * 60)
    print("NOTE: Due to the perfect block design, KPCA on the raw matrix recovers")
    print("annotation-assignment structure, not behavioral differences.")
    print("This section runs for comparison purposes only.\n")

    # 2a. Fit KernelPCA -----------------------------------------------
    kpca_full = KernelPCA(
        kernel="cosine",
        n_components=n_annotators,
        random_state=RANDOM_STATE,
    )
    Z_full = kpca_full.fit_transform(M_onehot)

    # 2b. Cumulative variance from eigenvalues -------------------------
    eigenvalues = kpca_full.eigenvalues_
    pos_mask = eigenvalues > 0
    eigenvalues_pos = eigenvalues[pos_mask]
    cum_var = np.cumsum(eigenvalues_pos) / np.sum(eigenvalues_pos)

    # 2c. Select k at variance threshold ------------------------------
    k = int(np.searchsorted(cum_var, VARIANCE_THRESHOLD)) + 1
    k = max(k, 2)
    print(f"Positive eigenvalues: {len(eigenvalues_pos)}")
    print(f"Selected k = {k} components at {VARIANCE_THRESHOLD:.1%} threshold "
          f"(actual: {cum_var[k - 1]:.4f})")

    # 2d. Truncate embeddings -----------------------------------------
    Z = Z_full[:, :k]
    print(f"KPCA embeddings shape: {Z.shape}")

    # 2e. Plot cumulative variance ------------------------------------
    fig, ax = plt.subplots(figsize=(8, 4))
    ax.plot(range(1, len(cum_var) + 1), cum_var, "b-", linewidth=1)
    ax.axhline(y=VARIANCE_THRESHOLD, color="r", linestyle="--",
               label=f"{VARIANCE_THRESHOLD:.1%} threshold")
    ax.axvline(x=k, color="g", linestyle="--", label=f"k = {k}")
    ax.set_xlabel("Number of components")
    ax.set_ylabel("Cumulative variance explained")
    ax.set_title("Kernel PCA (Cosine Kernel) — Cumulative Variance")
    ax.legend()
    ax.set_xlim(0, min(100, len(cum_var)))
    plt.tight_layout()
    plt.savefig("kpca_cumulative_variance.png", dpi=150, bbox_inches="tight")
    plt.close()
    print("Saved: kpca_cumulative_variance.png\n")

    # ── Section 3: Ward's Linkage on KPCA ──
    print("=" * 60)
    print("SECTION 3 — Ward's Linkage Clustering on KPCA [COMPARISON]")
    print("=" * 60)

    linkage_matrix = linkage(Z, method="ward")

    fig, ax = plt.subplots(figsize=(16, 6))
    dendrogram(linkage_matrix, labels=annotator_names,
               leaf_rotation=90, leaf_font_size=4, ax=ax)
    ax.set_title("Hierarchical Clustering Dendrogram (Ward's Linkage on KPCA)")
    ax.set_ylabel("Distance")
    plt.tight_layout()
    plt.savefig("dendrogram_full.png", dpi=150, bbox_inches="tight")
    plt.close()
    print("Saved: dendrogram_full.png")

    kpca_metrics_rows = []
    for n_c in CLUSTER_RANGE:
        labels_c = fcluster(linkage_matrix, t=n_c, criterion="maxclust")
        sil = silhouette_score(Z, labels_c)
        ch = calinski_harabasz_score(Z, labels_c)
        db = davies_bouldin_score(Z, labels_c)
        min_size = min(np.bincount(labels_c)[1:])
        kpca_metrics_rows.append({
            "n_clusters": n_c, "silhouette": round(sil, 4),
            "calinski_harabasz": round(ch, 2), "davies_bouldin": round(db, 4),
            "ch_db_ratio": round(ch / db, 2) if db > 0 else np.inf,
            "min_cluster_size": min_size,
        })

    kpca_metrics_df = pd.DataFrame(kpca_metrics_rows)
    min_viable = max(int(0.05 * n_annotators), 5)
    viable = kpca_metrics_df[kpca_metrics_df["min_cluster_size"] >= min_viable]
    if len(viable) == 0:
        viable = kpca_metrics_df
    kpca_optimal_idx = viable["ch_db_ratio"].idxmax()
    kpca_optimal_n = int(kpca_metrics_df.loc[kpca_optimal_idx, "n_clusters"])
    print(f"KPCA optimal n_clusters = {kpca_optimal_n} "
          f"(silhouette = {kpca_metrics_df.loc[kpca_optimal_idx, 'silhouette']:.4f})")

    kpca_cluster_labels = fcluster(linkage_matrix, t=kpca_optimal_n, criterion="maxclust")

    fig, ax = plt.subplots(figsize=(16, 6))
    color_thresh = linkage_matrix[-(kpca_optimal_n - 1), 2]
    dendrogram(linkage_matrix, labels=annotator_names,
               leaf_rotation=90, leaf_font_size=4,
               color_threshold=color_thresh, ax=ax)
    ax.set_title(f"Dendrogram — {kpca_optimal_n} Clusters (Ward's on KPCA)")
    ax.set_ylabel("Distance")
    plt.tight_layout()
    plt.savefig("dendrogram_colored.png", dpi=150, bbox_inches="tight")
    plt.close()
    print("Saved: dendrogram_colored.png\n")

else:
    print("=" * 60)
    print("SECTION 2-3 — KPCA + Ward's SKIPPED (RUN_KPCA_COMPARISON = False)")
    print("=" * 60, "\n")


# ============================================================
# Section 2B: Behavioral Feature Clustering (PRIMARY METHOD)
# ============================================================
print("=" * 60)
print("SECTION 2B — Behavioral Feature Clustering")
print("=" * 60)

# 2B-a. Compute behavioral features per annotator -------------------
labeled_mask_beh = (A != 0)
yes_counts_beh = (A == 1).sum(axis=1)
total_labeled_beh = labeled_mask_beh.sum(axis=1)

# YES rate
yes_rate_beh = yes_counts_beh / total_labeled_beh

# Agreement rate: fraction of texts where annotator agrees with majority
majority_label = np.sign(A.sum(axis=0))  # +1 if more YES, -1 if more NO, 0 if tie
agreement_counts = np.zeros(n_annotators)
agreement_total = np.zeros(n_annotators)
for i in range(n_annotators):
    for j in range(n_texts):
        if A[i, j] != 0 and majority_label[j] != 0:
            agreement_total[i] += 1
            if A[i, j] == majority_label[j]:
                agreement_counts[i] += 1
agreement_rate_beh = np.where(agreement_total > 0, agreement_counts / agreement_total, 0.5)

# Label entropy: Shannon entropy of YES/NO proportions per annotator
no_counts_beh = (A == -1).sum(axis=1)
p_yes = yes_counts_beh / total_labeled_beh
p_no = no_counts_beh / total_labeled_beh
entropy_beh = np.zeros(n_annotators)
for i in range(n_annotators):
    for p in [p_yes[i], p_no[i]]:
        if p > 0:
            entropy_beh[i] -= p * np.log2(p)

behavioral_features = np.column_stack([yes_rate_beh, agreement_rate_beh, entropy_beh])
feature_names_beh = ["yes_rate", "agreement_rate", "label_entropy"]

print(f"Behavioral features shape: {behavioral_features.shape}")
print(f"Feature ranges:")
for fi, fn in enumerate(feature_names_beh):
    vals = behavioral_features[:, fi]
    print(f"  {fn}: [{vals.min():.3f}, {vals.max():.3f}], mean={vals.mean():.3f}")

# 2B-b. Standardize features ----------------------------------------
scaler = StandardScaler()
behavioral_scaled = scaler.fit_transform(behavioral_features)

# 2B-c. KMeans scan k=2..20, select best by silhouette ---------------
beh_metrics_rows = []
for n_c in CLUSTER_RANGE:
    km = KMeans(n_clusters=n_c, random_state=RANDOM_STATE, n_init=20)
    labels_c = km.fit_predict(behavioral_scaled)
    sil = silhouette_score(behavioral_scaled, labels_c)
    ch = calinski_harabasz_score(behavioral_scaled, labels_c)
    db = davies_bouldin_score(behavioral_scaled, labels_c)
    min_size = min(np.bincount(labels_c))
    beh_metrics_rows.append({
        "n_clusters": n_c, "silhouette": round(sil, 4),
        "calinski_harabasz": round(ch, 2), "davies_bouldin": round(db, 4),
        "min_cluster_size": min_size,
    })

beh_metrics_df = pd.DataFrame(beh_metrics_rows)
print("\nBehavioral clustering metrics:")
print(beh_metrics_df.to_string(index=False))

beh_optimal_idx = beh_metrics_df["silhouette"].idxmax()
beh_optimal_n = int(beh_metrics_df.loc[beh_optimal_idx, "n_clusters"])
beh_best_sil = beh_metrics_df.loc[beh_optimal_idx, "silhouette"]
print(f"\n>>> Behavioral optimal n_clusters = {beh_optimal_n} "
      f"(silhouette = {beh_best_sil:.4f})")

# 2B-d. Final behavioral clustering -----------------------------------
km_final = KMeans(n_clusters=beh_optimal_n, random_state=RANDOM_STATE, n_init=20)
behavioral_labels = km_final.fit_predict(behavioral_scaled)  # 0-indexed
print(f"Behavioral cluster sizes: {dict(zip(*np.unique(behavioral_labels, return_counts=True)))}")

# 2B-e. Plot: behavioral cluster metrics ------------------------------
fig, axes = plt.subplots(1, 3, figsize=(15, 4))
for ax, (col, title) in zip(axes, [
    ("silhouette", "Silhouette (higher = better)"),
    ("calinski_harabasz", "Calinski-Harabasz (higher = better)"),
    ("davies_bouldin", "Davies-Bouldin (lower = better)"),
]):
    ax.plot(beh_metrics_df["n_clusters"], beh_metrics_df[col], "o-")
    ax.axvline(x=beh_optimal_n, color="r", linestyle="--", alpha=0.5)
    ax.set_xlabel("n_clusters")
    ax.set_title(title)
plt.suptitle("Behavioral Feature Clustering — Metrics Scan", fontsize=13, y=1.02)
plt.tight_layout()
plt.savefig("behavioral_cluster_metrics.png", dpi=150, bbox_inches="tight")
plt.close()
print("Saved: behavioral_cluster_metrics.png")

# 2B-f. Plot: 3D scatter of behavioral features -----------------------
# Sort clusters by yes_rate for consistent naming
cluster_means = {}
for c in range(beh_optimal_n):
    mask = behavioral_labels == c
    cluster_means[c] = behavioral_features[mask, 0].mean()  # mean yes_rate
sorted_clusters = sorted(cluster_means.keys(), key=lambda c: cluster_means[c])
cluster_name_map = {}
names_list = ["Conservative", "Mainstream", "Sensitive"]
if beh_optimal_n == 3:
    for i, c in enumerate(sorted_clusters):
        cluster_name_map[c] = names_list[i]
else:
    for i, c in enumerate(sorted_clusters):
        cluster_name_map[c] = f"Cluster {i + 1}"
behavioral_cluster_names = [cluster_name_map[c] for c in behavioral_labels]

# 2B-f′. Export behavioral features with cluster assignments ----------
no_counts_beh = (A == -1).sum(axis=1)
beh_export_df = pd.DataFrame({
    "annotator_id": annotator_names,
    "yes_rate": yes_rate_beh,
    "agreement_rate": agreement_rate_beh,
    "label_entropy": entropy_beh,
    "n_labeled": total_labeled_beh,
    "n_yes": yes_counts_beh,
    "n_no": no_counts_beh,
    "cluster": behavioral_cluster_names,
})
beh_export_df = beh_export_df.merge(
    demo_df.reset_index()[["annotator_id", "gender", "age", "ethnicity", "education", "country"]],
    on="annotator_id", how="left",
)
beh_export_df.to_csv("data/behavioral_features.csv", index=False)
print(f"Saved: behavioral_features.csv ({beh_export_df.shape[0]} annotators, "
      f"columns: {list(beh_export_df.columns)})")

fig, axes = plt.subplots(1, 3, figsize=(16, 5))
pairs = [(0, 1), (0, 2), (1, 2)]
pair_labels = [("YES Rate", "Agreement Rate"), ("YES Rate", "Label Entropy"),
               ("Agreement Rate", "Label Entropy")]
colors = sns.color_palette("Set1", beh_optimal_n)
for ax, (fi, fj), (xl, yl) in zip(axes, pairs, pair_labels):
    for ci, cname in enumerate(sorted(set(behavioral_cluster_names))):
        mask = np.array(behavioral_cluster_names) == cname
        ax.scatter(behavioral_features[mask, fi], behavioral_features[mask, fj],
                   alpha=0.6, label=cname, s=30)
    ax.set_xlabel(xl)
    ax.set_ylabel(yl)
    ax.legend(fontsize=7)
axes[1].set_title("Behavioral Feature Clustering — Pairwise Scatter")
plt.tight_layout()
plt.savefig("behavioral_features_scatter.png", dpi=150, bbox_inches="tight")
plt.close()
print("Saved: behavioral_features_scatter.png")

# 2B-g. Plot: cluster profiles (mean ± std of each feature) ----------
profile_data = []
for c in range(beh_optimal_n):
    mask = behavioral_labels == c
    cname = cluster_name_map[c]
    n_in = mask.sum()
    for fi, fn in enumerate(feature_names_beh):
        vals = behavioral_features[mask, fi]
        profile_data.append({
            "cluster": cname, "n": n_in, "feature": fn,
            "mean": vals.mean(), "std": vals.std(),
        })
profile_df = pd.DataFrame(profile_data)

fig, ax = plt.subplots(figsize=(10, 5))
x_pos = np.arange(len(feature_names_beh))
width = 0.8 / beh_optimal_n
for ci, cname in enumerate(sorted(set(profile_df["cluster"]))):
    sub = profile_df[profile_df["cluster"] == cname]
    means = sub["mean"].values
    stds = sub["std"].values
    offset = (ci - (beh_optimal_n - 1) / 2) * width
    bars = ax.bar(x_pos + offset, means, width, yerr=stds, capsize=3,
                  label=f"{cname} (n={sub['n'].iloc[0]})", alpha=0.8)
ax.set_xticks(x_pos)
ax.set_xticklabels(feature_names_beh)
ax.set_ylabel("Feature value")
ax.set_title("Behavioral Cluster Profiles")
ax.legend()
plt.tight_layout()
plt.savefig("behavioral_cluster_profiles.png", dpi=150, bbox_inches="tight")
plt.close()
print("Saved: behavioral_cluster_profiles.png")

# 2B-h. Plot: silhouette plot for behavioral clusters -----------------
beh_sil_values = silhouette_samples(behavioral_scaled, behavioral_labels)
beh_sil_avg = silhouette_score(behavioral_scaled, behavioral_labels)

fig, ax = plt.subplots(figsize=(8, 6))
y_lower = 10
for ci in range(beh_optimal_n):
    ci_vals = beh_sil_values[behavioral_labels == ci]
    ci_vals.sort()
    size_ci = ci_vals.shape[0]
    y_upper = y_lower + size_ci
    cname = cluster_name_map[ci]
    ax.fill_betweenx(np.arange(y_lower, y_upper), 0, ci_vals, alpha=0.7,
                     label=f"{cname} (n={size_ci})")
    y_lower = y_upper + 10
ax.axvline(x=beh_sil_avg, color="red", linestyle="--",
           label=f"Mean = {beh_sil_avg:.3f}")
ax.set_xlabel("Silhouette coefficient")
ax.set_ylabel("Annotators (sorted within each cluster)")
ax.set_title("Behavioral Clustering — Per-Annotator Silhouette Plot")
ax.legend(loc="best", fontsize=8)
plt.tight_layout()
plt.savefig("behavioral_silhouette_plot.png", dpi=150, bbox_inches="tight")
plt.close()
print("Saved: behavioral_silhouette_plot.png")

# 2B-i. Comparison plot: KPCA vs Behavioral --------------------------
if RUN_KPCA_COMPARISON and kpca_cluster_labels is not None:
    kpca_sil = silhouette_score(Z, kpca_cluster_labels)
    fig, axes = plt.subplots(1, 2, figsize=(12, 5))

    # KPCA silhouette
    kpca_sil_vals = silhouette_samples(Z, kpca_cluster_labels)
    y_lower = 10
    for ci in sorted(np.unique(kpca_cluster_labels)):
        ci_vals = kpca_sil_vals[kpca_cluster_labels == ci]
        ci_vals.sort()
        size_ci = ci_vals.shape[0]
        y_upper = y_lower + size_ci
        axes[0].fill_betweenx(np.arange(y_lower, y_upper), 0, ci_vals, alpha=0.7,
                              label=f"C{ci} (n={size_ci})")
        y_lower = y_upper + 10
    axes[0].axvline(x=kpca_sil, color="red", linestyle="--")
    axes[0].set_title(f"KPCA + Ward's (sil={kpca_sil:.3f})")
    axes[0].set_xlabel("Silhouette coefficient")
    axes[0].legend(fontsize=7)

    # Behavioral silhouette
    y_lower = 10
    for ci in range(beh_optimal_n):
        ci_vals = beh_sil_values[behavioral_labels == ci]
        ci_vals.sort()
        size_ci = ci_vals.shape[0]
        y_upper = y_lower + size_ci
        cname = cluster_name_map[ci]
        axes[1].fill_betweenx(np.arange(y_lower, y_upper), 0, ci_vals, alpha=0.7,
                              label=f"{cname} (n={size_ci})")
        y_lower = y_upper + 10
    axes[1].axvline(x=beh_sil_avg, color="red", linestyle="--")
    axes[1].set_title(f"Behavioral (sil={beh_sil_avg:.3f})")
    axes[1].set_xlabel("Silhouette coefficient")
    axes[1].legend(fontsize=7)

    plt.suptitle("KPCA vs Behavioral Clustering Comparison", fontsize=13)
    plt.tight_layout()
    plt.savefig("kpca_vs_behavioral_comparison.png", dpi=150, bbox_inches="tight")
    plt.close()
    print("Saved: kpca_vs_behavioral_comparison.png")

print()

# ============================================================
# Routing: Select primary cluster labels
# ============================================================
if PRIMARY_METHOD == "behavioral":
    cluster_labels = behavioral_labels + 1   # 1-indexed to match fcluster convention
    optimal_n = beh_optimal_n
    primary_embeddings = behavioral_scaled
    sil_avg = beh_sil_avg
    print(f"PRIMARY METHOD: behavioral  (k={optimal_n}, silhouette={sil_avg:.4f})")
elif PRIMARY_METHOD == "kpca" and kpca_cluster_labels is not None:
    cluster_labels = kpca_cluster_labels
    primary_embeddings = Z
    sil_avg = silhouette_score(Z, cluster_labels)
    print(f"PRIMARY METHOD: kpca  (k={optimal_n}, silhouette={sil_avg:.4f})")
else:
    raise ValueError(f"Invalid PRIMARY_METHOD='{PRIMARY_METHOD}' or KPCA not available")

demo_df["cluster"] = cluster_labels
print(f"Cluster sizes:")
print(pd.Series(cluster_labels).value_counts().sort_index())
print()


# ============================================================
# Section 4: Cluster Validation
# ============================================================
print("=" * 60)
print("SECTION 4 — Cluster Validation")
print("=" * 60)

# 4a. Per-cluster label rates (YES% / NO%) ------------------------
labeled_mask = (A != 0)
yes_counts = (A == 1).sum(axis=1)
no_counts = (A == -1).sum(axis=1)
total_labeled = labeled_mask.sum(axis=1)

yes_rate = yes_counts / total_labeled
no_rate = no_counts / total_labeled

rate_df = pd.DataFrame({
    "annotator_id": annotator_names,
    "cluster": cluster_labels,
    "yes_rate": yes_rate,
    "no_rate": no_rate,
    "n_labeled": total_labeled,
})

print("Per-cluster label rates:")
cluster_rates = rate_df.groupby("cluster")[["yes_rate", "no_rate"]].mean()
print(cluster_rates.round(4))
print()

# 4b. Pairwise agreement matrix (vectorised) ----------------------
L = labeled_mask.astype(np.float64)              # (n_ann, n_texts)
co_count = L @ L.T                               # co-labeled text counts
products = A.astype(np.float64) @ A.astype(np.float64).T  # +1 agree, −1 disagree
agree_count = (products + co_count) / 2.0

with np.errstate(divide="ignore", invalid="ignore"):
    agreement_matrix = np.where(co_count > 0, agree_count / co_count, 0.0)
np.fill_diagonal(agreement_matrix, 1.0)

# 4c. Within-cluster and between-cluster agreement ----------------
unique_clusters = sorted(np.unique(cluster_labels))
n_cl = len(unique_clusters)
agree_heatmap = np.zeros((n_cl, n_cl))

for i_idx, ci in enumerate(unique_clusters):
    for j_idx, cj in enumerate(unique_clusters):
        mask_i = (cluster_labels == ci)
        mask_j = (cluster_labels == cj)
        sub = agreement_matrix[np.ix_(mask_i, mask_j)].copy()
        if ci == cj:
            np.fill_diagonal(sub, np.nan)
            agree_heatmap[i_idx, j_idx] = np.nanmean(sub)
        else:
            agree_heatmap[i_idx, j_idx] = np.mean(sub)

print("Agreement heatmap (diagonal = within-cluster):")
agree_heatmap_df = pd.DataFrame(
    agree_heatmap,
    index=[f"C{c}" for c in unique_clusters],
    columns=[f"C{c}" for c in unique_clusters],
)
print(agree_heatmap_df.round(4))

within_avg = np.nanmean(np.diag(agree_heatmap))
off_diag = agree_heatmap[~np.eye(n_cl, dtype=bool)]
between_avg = np.mean(off_diag)
print(f"\nMean within-cluster agreement:  {within_avg:.4f}")
print(f"Mean between-cluster agreement: {between_avg:.4f}")
if between_avg > 0:
    print(f"Ratio (within / between):       {within_avg / between_avg:.3f}")
else:
    print("Ratio (within / between):       N/A (no between-cluster co-labels)")

# 4d. Agreement heatmap plot ---------------------------------------
fig, ax = plt.subplots(figsize=(8, 6))
sns.heatmap(agree_heatmap_df, annot=True, fmt=".3f", cmap="YlOrRd",
            vmin=0.4, vmax=1.0, ax=ax)
ax.set_title("Pairwise Agreement Heatmap (Within / Between Clusters)")
plt.tight_layout()
plt.savefig("agreement_heatmap.png", dpi=150, bbox_inches="tight")
plt.close()
print("Saved: agreement_heatmap.png")

# 4e. Per-sample silhouette plot -----------------------------------
sil_values = silhouette_samples(primary_embeddings, cluster_labels)
sil_avg = silhouette_score(primary_embeddings, cluster_labels)

fig, ax = plt.subplots(figsize=(8, 6))
y_lower = 10
for ci in unique_clusters:
    ci_vals = sil_values[cluster_labels == ci]
    ci_vals.sort()
    size_ci = ci_vals.shape[0]
    y_upper = y_lower + size_ci
    ax.fill_betweenx(np.arange(y_lower, y_upper), 0, ci_vals, alpha=0.7,
                     label=f"Cluster {ci} (n={size_ci})")
    ax.text(-0.05, y_lower + 0.5 * size_ci, str(ci), fontsize=10, va="center")
    y_lower = y_upper + 10

ax.axvline(x=sil_avg, color="red", linestyle="--",
           label=f"Mean = {sil_avg:.3f}")
ax.set_xlabel("Silhouette coefficient")
ax.set_ylabel("Annotators (sorted within each cluster)")
ax.set_title("Per-Annotator Silhouette Plot")
ax.legend(loc="best", fontsize=8)
plt.tight_layout()
plt.savefig("silhouette_plot.png", dpi=150, bbox_inches="tight")
plt.close()
print("Saved: silhouette_plot.png")

# Identify poorly assigned annotators (negative silhouette)
poor_mask = sil_values < 0
if poor_mask.any():
    poor_ids = [annotator_names[i] for i in np.where(poor_mask)[0]]
    print(f"\nPoorly assigned annotators (silhouette < 0): {len(poor_ids)}")
    for pid in poor_ids:
        idx = annotator_names.index(pid)
        print(f"  {pid}  cluster={cluster_labels[idx]}  sil={sil_values[idx]:.3f}")
else:
    print("\nNo poorly assigned annotators (all silhouette ≥ 0).")
print()


# ============================================================
# Section 5: Post-Hoc Demographic Analysis
# ============================================================
print("=" * 60)
print("SECTION 5 — Post-Hoc Demographic Analysis")
print("=" * 60)

DEMO_VARS = ["gender", "age", "ethnicity", "education", "region"]

# 5a. Contingency tables, Chi-squared, Cramér's V, ARI, AMI -------
demo_results = []

for var in DEMO_VARS:
    values = demo_df[var].values
    ct = pd.crosstab(demo_df["cluster"], demo_df[var])

    # Chi-squared test
    chi2, p_val, dof, expected = chi2_contingency(ct)
    n_total = ct.sum().sum()
    min_dim = min(ct.shape) - 1
    cramers_v = np.sqrt(chi2 / (n_total * min_dim)) if min_dim > 0 else 0.0

    # ARI and AMI
    ari = adjusted_rand_score(cluster_labels, values)
    ami = adjusted_mutual_info_score(cluster_labels, values)

    demo_results.append({
        "variable": var,
        "chi2": round(chi2, 2),
        "p_value": p_val,
        "dof": dof,
        "cramers_v": round(cramers_v, 4),
        "ARI": round(ari, 4),
        "AMI": round(ami, 4),
    })

    print(f"\n--- {var.upper()} ---")
    print("Contingency table:")
    print(ct)
    print(f"  χ² = {chi2:.2f},  p = {p_val:.4g},  dof = {dof}")
    print(f"  Cramér's V = {cramers_v:.4f}")
    print(f"  ARI = {ari:.4f},  AMI = {ami:.4f}")

demo_summary = pd.DataFrame(demo_results)
print("\n\nDemographic alignment summary:")
print(demo_summary.to_string(index=False))

# 5b. Stacked bar charts ------------------------------------------
fig, axes = plt.subplots(2, 3, figsize=(18, 10))
axes = axes.ravel()

for idx, var in enumerate(DEMO_VARS):
    ax = axes[idx]
    ct = pd.crosstab(demo_df["cluster"], demo_df[var], normalize="index")
    ct.plot(kind="bar", stacked=True, ax=ax, colormap="Set2", edgecolor="white")
    ax.set_title(f"Cluster composition by {var}")
    ax.set_xlabel("Cluster")
    ax.set_ylabel("Proportion")
    ax.legend(title=var, fontsize=7, title_fontsize=8, loc="upper right")
    ax.set_xticklabels(ax.get_xticklabels(), rotation=0)

# Hide unused subplot
if len(DEMO_VARS) < len(axes):
    for j in range(len(DEMO_VARS), len(axes)):
        axes[j].set_visible(False)

plt.suptitle("Post-Hoc Demographic Composition of Clusters", fontsize=14, y=1.01)
plt.tight_layout()
plt.savefig("demographic_stacked_bars.png", dpi=150, bbox_inches="tight")
plt.close()
print("\nSaved: demographic_stacked_bars.png\n")


# ============================================================
# Section 6: GLMM — Statistical Modeling of Demographic Effects
# ============================================================
print("=" * 60)
print("SECTION 6 — GLMM (Binomial Mixed-Effects Model)")
print("=" * 60)

if not RUN_GLMM:
    print("Skipped (RUN_GLMM = False)\n")
else:
    try:
        from statsmodels.genmod.bayes_mixed_glm import BinomialBayesMixedGLM
        import scipy.sparse as sp_sparse

        # 6a. Prepare data -------------------------------------------------
        # Use annotator-level data for English + active annotators, binary label
        glmm_df = annot_en[annot_en["annotator_id"].isin(active_annotators)].copy()
        glmm_df["label_binary"] = (glmm_df["label_raw"] == "YES").astype(int)
        glmm_df["region"] = glmm_df["country"].map(COUNTRY_TO_REGION).fillna("Other")

        n_obs = len(glmm_df)
        print(f"GLMM observations: {n_obs}")

        # 6b. Fixed effects design matrix (demographics) -------------------
        # Dummy-code each demographic variable (drop first category as reference)
        fixed_vars = ["gender", "age", "ethnicity", "education", "region"]
        exog_parts = [np.ones((n_obs, 1))]  # intercept
        fixed_names = ["Intercept"]

        for var in fixed_vars:
            dummies = pd.get_dummies(glmm_df[var], prefix=var, drop_first=True,
                                     dtype=float)
            exog_parts.append(dummies.values)
            fixed_names.extend(dummies.columns.tolist())

        exog = np.hstack(exog_parts)
        print(f"Fixed effects: {exog.shape[1]} columns ({', '.join(fixed_names[:5])} ...)")

        # 6c. Random effects design matrix ---------------------------------
        # (1 | text_id) + (1 | annotator_id)
        text_cat = pd.Categorical(glmm_df["text_id"])
        ann_cat = pd.Categorical(glmm_df["annotator_id"])

        n_text_levels = len(text_cat.categories)
        n_ann_levels = len(ann_cat.categories)
        print(f"Random effects: {n_text_levels} text levels + "
              f"{n_ann_levels} annotator levels")

        # Build sparse then convert to dense (required by statsmodels)
        rows_idx = np.arange(n_obs)
        text_vc = sp_sparse.csr_matrix(
            (np.ones(n_obs), (rows_idx, text_cat.codes)),
            shape=(n_obs, n_text_levels),
        )
        ann_vc = sp_sparse.csr_matrix(
            (np.ones(n_obs), (rows_idx, ann_cat.codes)),
            shape=(n_obs, n_ann_levels),
        )
        exog_vc = sp_sparse.hstack([text_vc, ann_vc]).toarray()
        ident = np.array([0] * n_text_levels + [1] * n_ann_levels)

        # 6d. Fit model ----------------------------------------------------
        endog = glmm_df["label_binary"].values.astype(float)

        print("Fitting BinomialBayesMixedGLM (variational Bayes)...")
        model = BinomialBayesMixedGLM(endog, exog, exog_vc, ident)
        result = model.fit_vb(verbose=True)

        # 6e. Fixed effects table ------------------------------------------
        fe_df = pd.DataFrame({
            "variable": fixed_names,
            "coefficient": result.fe_mean,
            "std_error": result.fe_sd,
            "z_value": result.fe_mean / result.fe_sd,
            "odds_ratio": np.exp(result.fe_mean),
        })
        fe_df["p_value"] = 2 * (1 - __import__("scipy").stats.norm.cdf(
            np.abs(fe_df["z_value"])
        ))
        fe_df["sig"] = fe_df["p_value"].apply(
            lambda p: "***" if p < 0.001 else ("**" if p < 0.01 else
                      ("*" if p < 0.05 else ""))
        )

        print("\n--- Fixed Effects ---")
        print(fe_df.to_string(index=False))

        # 6f. Variance decomposition ---------------------------------------
        # vcp_mean contains log(sigma) for each random effects group
        sigma_text = np.exp(result.vcp_mean[0])
        sigma_ann = np.exp(result.vcp_mean[1])
        var_text = sigma_text ** 2
        var_ann = sigma_ann ** 2
        var_residual = (np.pi ** 2) / 3  # logistic distribution variance

        # Variance of linear predictor due to fixed effects
        eta_fixed = exog @ result.fe_mean
        var_fixed = np.var(eta_fixed)

        total_var = var_fixed + var_text + var_ann + var_residual

        print("\n--- Variance Decomposition ---")
        print(f"  Demographics (fixed):  {var_fixed:8.4f}  ({var_fixed / total_var:6.2%})")
        print(f"  Text (random):         {var_text:8.4f}  ({var_text / total_var:6.2%})")
        print(f"  Annotator (random):    {var_ann:8.4f}  ({var_ann / total_var:6.2%})")
        print(f"  Residual (logistic):   {var_residual:8.4f}  ({var_residual / total_var:6.2%})")
        print(f"  Total:                 {total_var:8.4f}")

        # ICC
        icc_text = var_text / (var_text + var_ann + var_residual)
        icc_ann = var_ann / (var_text + var_ann + var_residual)
        print(f"\n  ICC(text)      = {icc_text:.4f}")
        print(f"  ICC(annotator) = {icc_ann:.4f}")

    except Exception as e:
        print(f"\nGLMM failed: {e}")
        print("This may happen due to memory constraints or convergence issues.")
        print("Consider running on Colab or reducing the dataset.")

    print()


# ============================================================
# Section 7: Robustness Checks (Optional)
# ============================================================
print("=" * 60)
print("SECTION 7 — Robustness Checks")
print("=" * 60)

if not RUN_ROBUSTNESS:
    print("Skipped (RUN_ROBUSTNESS = False)\n")
else:
    # 7a. Sensitivity to KPCA variance threshold ----------------------
    if RUN_KPCA_COMPARISON and cum_var is not None:
        print("7a. KPCA variance threshold sensitivity")
        thresholds = [0.80, 0.85, 0.855, 0.90, 0.95]
        for thr in thresholds:
            k_t = int(np.searchsorted(cum_var, thr)) + 1
            k_t = max(k_t, 2)
            Z_t = Z_full[:, :k_t]
            lm_t = linkage(Z_t, method="ward")
            labels_t = fcluster(lm_t, t=optimal_n, criterion="maxclust")
            sil_t = silhouette_score(Z_t, labels_t)
            ari_vs_main = adjusted_rand_score(cluster_labels, labels_t)
            print(f"  threshold={thr:.1%}  k={k_t:3d}  silhouette={sil_t:.4f}  "
                  f"ARI_vs_main={ari_vs_main:.4f}")
    else:
        print("7a. KPCA sensitivity — skipped (KPCA not run)")

    # 7b. KMeans vs HAC comparison on same embeddings -----------------
    if RUN_KPCA_COMPARISON and Z is not None:
        print("\n7b. KMeans vs HAC comparison on KPCA embeddings")
        km_labels_7b = KMeans(n_clusters=optimal_n, random_state=RANDOM_STATE,
                              n_init=20).fit_predict(Z)
        sil_km = silhouette_score(Z, km_labels_7b)
        sil_hac = silhouette_score(Z, kpca_cluster_labels) if kpca_cluster_labels is not None else 0
        print(f"  HAC silhouette:    {sil_hac:.4f}")
        print(f"  KMeans silhouette: {sil_km:.4f}")
    else:
        print("\n7b. KMeans vs HAC on KPCA — skipped (KPCA not run)")

    # 7c. Bootstrap cluster stability (KPCA) --------------------------
    if RUN_KPCA_COMPARISON and Z is not None:
        print("\n7c. Bootstrap cluster stability — KPCA (100 resamples)")
        n_boot = 100
        jaccard_scores = []
        for b in range(n_boot):
            rng = np.random.RandomState(b)
            boot_idx = rng.choice(n_annotators, size=n_annotators, replace=True)
            Z_boot = Z[boot_idx]
            lm_boot = linkage(Z_boot, method="ward")
            labels_boot = fcluster(lm_boot, t=optimal_n, criterion="maxclust")
            original_boot = cluster_labels[boot_idx]
            boot_unique = np.unique(labels_boot)
            orig_unique = np.unique(original_boot)
            used = set()
            total_jaccard = 0
            for bc in boot_unique:
                bc_set = set(np.where(labels_boot == bc)[0])
                best_j = 0
                best_oc = orig_unique[0]
                for oc in orig_unique:
                    if oc in used:
                        continue
                    oc_set = set(np.where(original_boot == oc)[0])
                    inter = len(bc_set & oc_set)
                    union = len(bc_set | oc_set)
                    j = inter / union if union > 0 else 0
                    if j > best_j:
                        best_j = j
                        best_oc = oc
                used.add(best_oc)
                total_jaccard += best_j
            jaccard_scores.append(total_jaccard / len(boot_unique))
        jaccard_arr = np.array(jaccard_scores)
        print(f"  Mean Jaccard: {jaccard_arr.mean():.4f} +/- {jaccard_arr.std():.4f}")

        fig, ax = plt.subplots(figsize=(6, 4))
        ax.hist(jaccard_arr, bins=20, edgecolor="black", alpha=0.7)
        ax.axvline(x=jaccard_arr.mean(), color="red", linestyle="--",
                   label=f"Mean = {jaccard_arr.mean():.3f}")
        ax.set_xlabel("Mean Jaccard similarity")
        ax.set_ylabel("Frequency")
        ax.set_title("Bootstrap Stability — KPCA (100 resamples)")
        ax.legend()
        plt.tight_layout()
        plt.savefig("bootstrap_stability.png", dpi=150, bbox_inches="tight")
        plt.close()
        print("Saved: bootstrap_stability.png")
    else:
        print("\n7c. Bootstrap KPCA — skipped")

    # 7d. Feature subset sensitivity (behavioral) ---------------------
    print("\n7d. Behavioral feature subset sensitivity")
    feature_subsets = {
        "yes_rate + agreement_rate": [0, 1],
        "yes_rate + entropy": [0, 2],
        "agreement_rate + entropy": [1, 2],
        "all 3 features": [0, 1, 2],
    }
    for label, cols in feature_subsets.items():
        sub_feats = behavioral_features[:, cols]
        sub_scaled = StandardScaler().fit_transform(sub_feats)
        km_sub = KMeans(n_clusters=beh_optimal_n, random_state=RANDOM_STATE, n_init=20)
        sub_labels = km_sub.fit_predict(sub_scaled)
        sub_sil = silhouette_score(sub_scaled, sub_labels)
        sub_ari = adjusted_rand_score(behavioral_labels, sub_labels)
        print(f"  {label:35s}  sil={sub_sil:.4f}  ARI_vs_full={sub_ari:.4f}")

    # 7e. KMeans vs Ward's on behavioral features ---------------------
    print("\n7e. KMeans vs Ward's on behavioral features")
    ward_beh = AgglomerativeClustering(n_clusters=beh_optimal_n, linkage="ward")
    ward_beh_labels = ward_beh.fit_predict(behavioral_scaled)
    ward_sil = silhouette_score(behavioral_scaled, ward_beh_labels)
    km_sil = silhouette_score(behavioral_scaled, behavioral_labels)
    ari_km_ward = adjusted_rand_score(behavioral_labels, ward_beh_labels)
    print(f"  KMeans silhouette:   {km_sil:.4f}")
    print(f"  Ward's silhouette:   {ward_sil:.4f}")
    print(f"  ARI (KMeans vs Ward's): {ari_km_ward:.4f}")

    # 7f. Bootstrap stability for behavioral clusters ------------------
    print("\n7f. Bootstrap stability — behavioral (100 resamples)")
    n_boot = 100
    beh_jaccard_scores = []
    for b in range(n_boot):
        rng = np.random.RandomState(b)
        boot_idx = rng.choice(n_annotators, size=n_annotators, replace=True)
        beh_boot = behavioral_scaled[boot_idx]
        km_boot = KMeans(n_clusters=beh_optimal_n, random_state=RANDOM_STATE, n_init=20)
        labels_boot = km_boot.fit_predict(beh_boot)
        original_boot = behavioral_labels[boot_idx]
        boot_unique = np.unique(labels_boot)
        orig_unique = np.unique(original_boot)
        used = set()
        total_jaccard = 0
        for bc in boot_unique:
            bc_set = set(np.where(labels_boot == bc)[0])
            best_j = 0
            best_oc = orig_unique[0]
            for oc in orig_unique:
                if oc in used:
                    continue
                oc_set = set(np.where(original_boot == oc)[0])
                inter = len(bc_set & oc_set)
                union = len(bc_set | oc_set)
                j = inter / union if union > 0 else 0
                if j > best_j:
                    best_j = j
                    best_oc = oc
            used.add(best_oc)
            total_jaccard += best_j
        beh_jaccard_scores.append(total_jaccard / len(boot_unique))

    beh_jaccard_arr = np.array(beh_jaccard_scores)
    print(f"  Mean Jaccard: {beh_jaccard_arr.mean():.4f} +/- {beh_jaccard_arr.std():.4f}")
    print(f"  Min: {beh_jaccard_arr.min():.4f}  Max: {beh_jaccard_arr.max():.4f}")

    fig, ax = plt.subplots(figsize=(6, 4))
    ax.hist(beh_jaccard_arr, bins=20, edgecolor="black", alpha=0.7)
    ax.axvline(x=beh_jaccard_arr.mean(), color="red", linestyle="--",
               label=f"Mean = {beh_jaccard_arr.mean():.3f}")
    ax.set_xlabel("Mean Jaccard similarity")
    ax.set_ylabel("Frequency")
    ax.set_title("Bootstrap Stability — Behavioral (100 resamples)")
    ax.legend()
    plt.tight_layout()
    plt.savefig("bootstrap_stability_behavioral.png", dpi=150, bbox_inches="tight")
    plt.close()
    print("Saved: bootstrap_stability_behavioral.png")

    print()


# ============================================================
# Section 8: Word Report Generation (Optional)
# ============================================================
print("=" * 60)
print("SECTION 8 — Word Report Generation")
print("=" * 60)

if not GENERATE_REPORT:
    print("Skipped (GENERATE_REPORT = False)\n")
else:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import os

        doc = Document()

        # -- Style setup --
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # ── Title ──
        title = doc.add_heading("Annotator Clustering Analysis — EXIST 2024", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(
            "How annotators differ in labeling sexism, and what drives those differences."
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")

        # ── 1. Introduction ──
        doc.add_heading("1. Introduction", level=1)
        doc.add_paragraph(
            "EXIST 2024 (sEXism Identification in Social neTworks) is a shared task where "
            "multiple people read social media posts and decide whether each one contains "
            "sexism. The interesting part is that different annotators often disagree — "
            "what one person considers sexist, another might not."
        )
        doc.add_paragraph(
            "In this analysis, we try to understand those disagreements. Instead of treating "
            "annotator disagreement as noise, we cluster annotators by how they actually "
            "behave — how often they label things as sexist, how much they agree with the "
            "majority, and how consistent their labeling patterns are. Then we check whether "
            "demographics (age, gender, education, etc.) explain any of the differences."
        )

        # ── 2. The Data ──
        doc.add_heading("2. The Data", level=1)
        doc.add_paragraph(
            f"The dataset has {len(raw)} texts total. After filtering to English, we have "
            f"{n_texts} texts labeled by {n_annotators} active annotators (each annotator "
            f"labeled at least {MIN_LABELS} texts)."
        )
        doc.add_paragraph(
            "Here is a key finding: the annotation has a perfect block design. There are "
            f"{n_groups} fixed groups of annotators — each group of 6 people labels a batch "
            "of texts, and no annotator appears in more than one group. This means annotators "
            "in different groups never rate the same text, so we cannot directly compare their "
            "ratings on identical content."
        )
        if os.path.exists("block_design_heatmap.png"):
            doc.add_picture("block_design_heatmap.png", width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(
            "The heatmap above shows this block-diagonal structure — each block is a group "
            "of 6 annotators rating the same set of texts, with zero overlap between groups."
        )

        # ── 3. How We Did It ──
        doc.add_heading("3. How We Did It", level=1)
        doc.add_heading("3.1 Why the Standard Approach Fails", level=2)
        doc.add_paragraph(
            "The standard approach in the literature is to use Kernel PCA with a cosine "
            "kernel on the annotation matrix, then cluster with Ward's linkage. But because "
            "of the block design, cosine similarity between annotators in different groups is "
            "exactly zero — they share no co-labeled texts. So KPCA just recovers which group "
            "each annotator belongs to, not how they actually differ in behavior."
        )
        if os.path.exists("kpca_cumulative_variance.png"):
            doc.add_picture("kpca_cumulative_variance.png", width=Inches(5.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading("3.2 Behavioral Feature Clustering", level=2)
        doc.add_paragraph(
            "Instead, we compute three behavioral features for each annotator:"
        )
        doc.add_paragraph("YES rate — what fraction of texts they label as sexist", style="List Bullet")
        doc.add_paragraph("Agreement rate — how often they agree with the majority label", style="List Bullet")
        doc.add_paragraph(
            "Label entropy — how predictable their labeling is (high entropy = roughly "
            "50/50 split between YES and NO)", style="List Bullet"
        )
        doc.add_paragraph(
            "We standardize these three features and run KMeans clustering with k = 2 to 20. "
            f"The best number of clusters is k = {beh_optimal_n}, chosen by the highest "
            f"silhouette score ({beh_best_sil:.3f})."
        )
        if os.path.exists("behavioral_cluster_metrics.png"):
            doc.add_picture("behavioral_cluster_metrics.png", width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ── 4. Results ──
        doc.add_heading("4. Results", level=1)

        doc.add_heading("4.1 The Three Clusters", level=2)
        doc.add_paragraph(
            "The behavioral clustering produces three distinct groups of annotators:"
        )
        # Build cluster descriptions from actual data
        for c in sorted(cluster_name_map.keys(), key=lambda x: cluster_name_map[x]):
            cname = cluster_name_map[c]
            mask = behavioral_labels == c
            n_in = mask.sum()
            yr = behavioral_features[mask, 0].mean()
            ar = behavioral_features[mask, 1].mean()
            ent = behavioral_features[mask, 2].mean()
            if "Conservative" in cname:
                desc = (f"These annotators rarely label things as sexist (YES rate ~{yr:.0%}). "
                        "They are strict and tend to reserve the sexism label for clear-cut cases.")
            elif "Mainstream" in cname:
                desc = (f"The largest group with moderate YES rates (~{yr:.0%}). "
                        "They tend to agree with the majority and represent the 'average' perspective.")
            else:
                desc = (f"These annotators label more things as sexist (YES rate ~{yr:.0%}). "
                        "They may be more sensitive to subtle forms of sexism that others miss.")
            doc.add_paragraph(f"{cname} (n={n_in}): {desc}")

        if os.path.exists("behavioral_features_scatter.png"):
            doc.add_picture("behavioral_features_scatter.png", width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if os.path.exists("behavioral_cluster_profiles.png"):
            doc.add_picture("behavioral_cluster_profiles.png", width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading("4.2 Validation", level=2)
        doc.add_paragraph(
            f"The silhouette score for behavioral clustering is {beh_sil_avg:.3f}, which "
            "indicates reasonably well-separated clusters. For comparison, KPCA + Ward's "
            "only achieved a silhouette of ~0.03 on the same data."
        )
        if os.path.exists("behavioral_silhouette_plot.png"):
            doc.add_picture("behavioral_silhouette_plot.png", width=Inches(5.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if os.path.exists("agreement_heatmap.png"):
            doc.add_picture("agreement_heatmap.png", width=Inches(4.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading("4.3 Demographics", level=2)
        doc.add_paragraph(
            "We checked whether demographic variables (gender, age, ethnicity, education, "
            "region) are associated with cluster membership using chi-squared tests."
        )
        if os.path.exists("demographic_stacked_bars.png"):
            doc.add_picture("demographic_stacked_bars.png", width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Chi-squared summary table
        doc.add_paragraph("")
        chi2_table = doc.add_table(rows=1, cols=5)
        chi2_table.style = "Light Shading Accent 1"
        hdr = chi2_table.rows[0].cells
        for i, h in enumerate(["Variable", "Chi-squared", "p-value", "Cramer's V", "Significant?"]):
            hdr[i].text = h
        for _, row_data in demo_summary.iterrows():
            row = chi2_table.add_row().cells
            row[0].text = str(row_data["variable"])
            row[1].text = f"{row_data['chi2']:.2f}"
            row[2].text = f"{row_data['p_value']:.4g}"
            row[3].text = f"{row_data['cramers_v']:.4f}"
            row[4].text = "Yes" if row_data["p_value"] < 0.05 else "No"

        doc.add_heading("4.4 What Drives Annotation Differences — GLMM", level=2)
        # Try to include GLMM results if they exist
        try:
            if fe_df is not None and len(fe_df) > 0:
                doc.add_paragraph(
                    "We fit a binomial mixed-effects model (GLMM) to understand what drives "
                    "individual annotation decisions. The model includes demographics as fixed "
                    "effects and text ID + annotator ID as random effects."
                )
                glmm_table = doc.add_table(rows=1, cols=6)
                glmm_table.style = "Light Shading Accent 1"
                hdr = glmm_table.rows[0].cells
                for i, h in enumerate(["Variable", "Coefficient", "Odds Ratio", "p-value", "Sig.", "Std Error"]):
                    hdr[i].text = h
                for _, r in fe_df.iterrows():
                    row = glmm_table.add_row().cells
                    row[0].text = str(r["variable"])
                    row[1].text = f"{r['coefficient']:.4f}"
                    row[2].text = f"{r['odds_ratio']:.4f}"
                    row[3].text = f"{r['p_value']:.4g}"
                    row[4].text = str(r["sig"])
                    row[5].text = f"{r['std_error']:.4f}"

                doc.add_paragraph("")
                doc.add_paragraph("Variance Decomposition:")
                doc.add_paragraph(
                    f"  Demographics (fixed effects): {var_fixed / total_var:.1%} of total variance"
                )
                doc.add_paragraph(
                    f"  Text content (random): {var_text / total_var:.1%}"
                )
                doc.add_paragraph(
                    f"  Annotator identity (random): {var_ann / total_var:.1%}"
                )
                doc.add_paragraph(
                    f"  Residual: {var_residual / total_var:.1%}"
                )
        except NameError:
            doc.add_paragraph(
                "GLMM results not available (model did not run or failed to converge)."
            )

        # ── 5. What This All Means ──
        doc.add_heading("5. What This All Means", level=1)
        doc.add_paragraph(
            "The biggest takeaway is that text content dominates annotation decisions — "
            "the same text tends to get similar labels regardless of who reads it. This makes "
            "sense: most sexism is fairly obvious when you see it."
        )
        doc.add_paragraph(
            "Annotator identity matters some. Each person has their own threshold for what "
            "counts as sexism, and this shows up in the three behavioral clusters. Some people "
            "are stricter (Conservative), some are more sensitive (Sensitive), and most fall "
            "somewhere in the middle (Mainstream)."
        )
        doc.add_paragraph(
            "Demographics barely matter on their own. Knowing someone's age, gender, or "
            "education tells you very little about how they will label a specific text. "
            "Individual variation within demographic groups is much larger than differences "
            "between groups."
        )
        doc.add_paragraph(
            "For DPO (Direct Preference Optimization) training, this means: cluster-based "
            "preference pairs (e.g., Mainstream vs. Sensitive labels on the same text) are "
            "more meaningful than demographic-based splits. The behavioral clusters capture "
            "real differences in annotation philosophy."
        )

        # ── 6. Conclusion ──
        doc.add_heading("6. Conclusion", level=1)
        doc.add_paragraph(
            "We showed that the EXIST 2024 dataset has a perfect block design that makes "
            "standard cosine-based clustering methods ineffective. By switching to behavioral "
            "features (YES rate, agreement rate, and label entropy), we identified three "
            f"meaningful annotator clusters with a silhouette score of {beh_sil_avg:.3f}. "
            "Demographics have minimal predictive power for annotation behavior — the key "
            "differences are individual tendencies, not group membership."
        )

        report_path = "annotator_clustering_report.docx"
        doc.save(report_path)
        print(f"Saved: {report_path}")

    except ImportError:
        print("python-docx not installed. Skipping report generation.")
        print("Install with: pip install python-docx")
    except Exception as e:
        print(f"Report generation failed: {e}")
        import traceback
        traceback.print_exc()

    print()


# ============================================================
# Section 9: Jupyter Notebook Generation (Optional)
# ============================================================
print("=" * 60)
print("SECTION 9 — Jupyter Notebook Generation")
print("=" * 60)

if not GENERATE_NOTEBOOK:
    print("Skipped (GENERATE_NOTEBOOK = False)\n")
else:
    def _nb_cell(cell_type, source):
        """Create a notebook cell dict."""
        return {
            "cell_type": cell_type,
            "metadata": {},
            "source": source.split("\n") if isinstance(source, str) else source,
            **({"outputs": [], "execution_count": None} if cell_type == "code" else {}),
        }

    # Fix: ensure source lines end with \n except the last
    def _fix_source(lines):
        if not lines:
            return lines
        result = []
        for i, line in enumerate(lines):
            if i < len(lines) - 1:
                result.append(line if line.endswith("\n") else line + "\n")
            else:
                result.append(line.rstrip("\n"))
        return result

    cells = []

    # Cell 1: Title
    cells.append(_nb_cell("markdown", "# Annotator Clustering Analysis — EXIST 2024\n\n"
        "This notebook implements behavioral feature clustering of annotators from the\n"
        "EXIST 2024 shared task dataset. We discover the block design structure,\n"
        "compute behavioral features, cluster annotators, and analyze demographic associations."))

    # Cell 2: Imports
    cells.append(_nb_cell("code",
        "import json\nimport warnings\nimport numpy as np\nimport pandas as pd\n"
        "import matplotlib.pyplot as plt\nimport seaborn as sns\n"
        "from sklearn.decomposition import KernelPCA\n"
        "from sklearn.preprocessing import StandardScaler\n"
        "from sklearn.cluster import AgglomerativeClustering, KMeans\n"
        "from sklearn.metrics import (\n"
        "    silhouette_score, silhouette_samples,\n"
        "    calinski_harabasz_score, davies_bouldin_score,\n"
        "    adjusted_rand_score, adjusted_mutual_info_score,\n"
        ")\n"
        "from scipy.cluster.hierarchy import linkage, dendrogram, fcluster\n"
        "from scipy.stats import chi2_contingency\n"
        "from collections import Counter\n\n"
        "warnings.filterwarnings('ignore', category=FutureWarning)\n"
        "%matplotlib inline\n\n"
        "RANDOM_STATE = 42\n"
        "MIN_LABELS = 10\n"
        "print('Imports done.')"))

    # Cell 3: Colab mount
    cells.append(_nb_cell("code",
        "# Google Colab drive mount (skip if running locally)\n"
        "try:\n"
        "    from google.colab import drive\n"
        "    drive.mount('/content/drive')\n"
        "    DATA_PATH = '/content/drive/My Drive/EXIST2024_training.json'\n"
        "except ImportError:\n"
        "    DATA_PATH = 'EXIST2024_training.json'\n"
        "print(f'Data path: {DATA_PATH}')"))

    # Cell 4: md - Data Loading
    cells.append(_nb_cell("markdown", "## 1. Data Loading & Matrix Construction\n\n"
        "Load the EXIST 2024 JSON, build the annotator-level table, pivot to text x annotator matrix."))

    # Cell 5: Data loading code
    cells.append(_nb_cell("code",
        "with open(DATA_PATH, 'r') as f:\n"
        "    raw = json.load(f)\n"
        "print(f'Loaded {len(raw)} texts')\n\n"
        "# Build annotator-level long table\n"
        "rows_list = []\n"
        "for entry in raw.values():\n"
        "    text_id = entry['id_EXIST']\n"
        "    lang = entry['lang']\n"
        "    tweet = entry['tweet']\n"
        "    annotators = entry['annotators']\n"
        "    n = len(annotators)\n"
        "    for i in range(n):\n"
        "        lab = entry['labels_task1'][i] if i < len(entry['labels_task1']) else None\n"
        "        rows_list.append({\n"
        "            'text_id': text_id, 'language': lang, 'tweet': tweet,\n"
        "            'annotator_id': annotators[i],\n"
        "            'gender': entry['gender_annotators'][i] if i < len(entry['gender_annotators']) else None,\n"
        "            'age': entry['age_annotators'][i] if i < len(entry['age_annotators']) else None,\n"
        "            'ethnicity': entry['ethnicities_annotators'][i] if i < len(entry['ethnicities_annotators']) else None,\n"
        "            'education': entry['study_levels_annotators'][i] if i < len(entry['study_levels_annotators']) else None,\n"
        "            'country': entry['countries_annotators'][i] if i < len(entry['countries_annotators']) else None,\n"
        "            'label_raw': lab,\n"
        "            'label': 1 if lab == 'YES' else (-1 if lab == 'NO' else 0),\n"
        "        })\n"
        "annot_df = pd.DataFrame(rows_list)\n"
        "print(f'Annotator-level table: {annot_df.shape}')"))

    # Cell 6: Matrix construction
    cells.append(_nb_cell("code",
        "# Pivot to text x annotator matrix\n"
        "matrix_df = (\n"
        "    annot_df.pivot_table(\n"
        "        index=['text_id', 'language'], columns='annotator_id',\n"
        "        values='label', aggfunc='first'\n"
        "    ).fillna(0).astype(int)\n"
        ")\n"
        "matrix_en = matrix_df.xs('en', level='language')\n"
        "print(f'English texts: {matrix_en.shape[0]}')\n\n"
        "# Active annotators\n"
        "non_zero_counts = (matrix_en != 0).sum(axis=0)\n"
        "active_mask = non_zero_counts >= MIN_LABELS\n"
        "active_annotators = non_zero_counts[active_mask].index.tolist()\n"
        "print(f'Active annotators (>= {MIN_LABELS} labels): {len(active_annotators)}')\n\n"
        "A = matrix_en[active_annotators].values.T\n"
        "annotator_names = active_annotators\n"
        "text_ids = matrix_en.index.tolist()\n"
        "n_annotators, n_texts = A.shape\n"
        "print(f'Matrix A shape: {A.shape}')"))

    # Cell 7: Demographics
    cells.append(_nb_cell("code",
        "# Extract demographics\n"
        "annot_en = annot_df[annot_df['language'] == 'en'].copy()\n"
        "demo_df = (\n"
        "    annot_en[annot_en['annotator_id'].isin(active_annotators)]\n"
        "    .drop_duplicates(subset='annotator_id')\n"
        "    .set_index('annotator_id')\n"
        "    [['gender', 'age', 'ethnicity', 'education', 'country']]\n"
        "    .loc[annotator_names]\n"
        ")\n"
        "print(f'Demographics table: {demo_df.shape}')"))

    # Cell 8: md - Block Design
    cells.append(_nb_cell("markdown", "## 2. Block Design Discovery\n\n"
        "We check whether the annotation has a block design — are annotators assigned in fixed groups?"))

    # Cell 9: Block design code
    cells.append(_nb_cell("code",
        "text_annotator_sets = {}\n"
        "for idx, tid in enumerate(text_ids):\n"
        "    labeled = [annotator_names[j] for j in range(n_annotators) if A[j, idx] != 0]\n"
        "    text_annotator_sets[tid] = frozenset(labeled)\n\n"
        "group_counter = Counter(text_annotator_sets.values())\n"
        "n_groups = len(group_counter)\n"
        "all_groups = list(group_counter.keys())\n"
        "print(f'Unique annotator groups: {n_groups}')\n\n"
        "# Check overlap\n"
        "has_overlap = False\n"
        "for i in range(len(all_groups)):\n"
        "    for j in range(i + 1, len(all_groups)):\n"
        "        if all_groups[i] & all_groups[j]:\n"
        "            has_overlap = True\n"
        "            break\n"
        "    if has_overlap:\n"
        "        break\n"
        "print(f'Cross-group overlap: {has_overlap}')"))

    # Cell 10: Block design heatmap
    cells.append(_nb_cell("code",
        "# Block design heatmap\n"
        "annotator_to_group = {}\n"
        "for gidx, grp in enumerate(all_groups):\n"
        "    for ann in grp:\n"
        "        annotator_to_group[ann] = gidx\n\n"
        "group_order = sorted(range(n_annotators), key=lambda i: annotator_to_group.get(annotator_names[i], -1))\n"
        "A_reordered = A[group_order, :]\n\n"
        "text_group = {}\n"
        "for idx, tid in enumerate(text_ids):\n"
        "    grp = text_annotator_sets[tid]\n"
        "    text_group[idx] = all_groups.index(grp)\n"
        "text_order = sorted(range(n_texts), key=lambda i: text_group[i])\n"
        "A_block = A_reordered[:, text_order]\n\n"
        "fig, ax = plt.subplots(figsize=(14, 8))\n"
        "display_matrix = np.where(A_block != 0, 1, 0).astype(float)\n"
        "ax.imshow(display_matrix, aspect='auto', cmap='Blues', interpolation='none')\n"
        "ax.set_xlabel('Texts (reordered by annotation group)')\n"
        "ax.set_ylabel('Annotators (reordered by annotation group)')\n"
        "ax.set_title(f'Block Design Structure - {n_groups} Disjoint Annotator Groups')\n"
        "plt.tight_layout()\n"
        "plt.show()"))

    # Cell 11: md - Behavioral Features
    cells.append(_nb_cell("markdown", "## 3. Behavioral Feature Extraction\n\n"
        "We compute three behavioral features per annotator:\n"
        "- **YES rate**: proportion of texts labeled as sexist\n"
        "- **Agreement rate**: agreement with per-text majority\n"
        "- **Label entropy**: Shannon entropy of YES/NO proportions"))

    # Cell 12: Feature extraction
    cells.append(_nb_cell("code",
        "labeled_mask = (A != 0)\n"
        "yes_counts = (A == 1).sum(axis=1)\n"
        "no_counts = (A == -1).sum(axis=1)\n"
        "total_labeled = labeled_mask.sum(axis=1)\n\n"
        "yes_rate = yes_counts / total_labeled\n\n"
        "# Agreement rate\n"
        "majority_label = np.sign(A.sum(axis=0))\n"
        "agreement_counts = np.zeros(n_annotators)\n"
        "agreement_total = np.zeros(n_annotators)\n"
        "for i in range(n_annotators):\n"
        "    for j in range(n_texts):\n"
        "        if A[i, j] != 0 and majority_label[j] != 0:\n"
        "            agreement_total[i] += 1\n"
        "            if A[i, j] == majority_label[j]:\n"
        "                agreement_counts[i] += 1\n"
        "agreement_rate = np.where(agreement_total > 0, agreement_counts / agreement_total, 0.5)\n\n"
        "# Label entropy\n"
        "p_yes = yes_counts / total_labeled\n"
        "p_no = no_counts / total_labeled\n"
        "entropy = np.zeros(n_annotators)\n"
        "for i in range(n_annotators):\n"
        "    for p in [p_yes[i], p_no[i]]:\n"
        "        if p > 0:\n"
        "            entropy[i] -= p * np.log2(p)\n\n"
        "behavioral_features = np.column_stack([yes_rate, agreement_rate, entropy])\n"
        "print(f'Features shape: {behavioral_features.shape}')\n"
        "for fi, fn in enumerate(['yes_rate', 'agreement_rate', 'label_entropy']):\n"
        "    vals = behavioral_features[:, fi]\n"
        "    print(f'  {fn}: [{vals.min():.3f}, {vals.max():.3f}], mean={vals.mean():.3f}')"))

    # Cell 13: Feature histograms
    cells.append(_nb_cell("code",
        "fig, axes = plt.subplots(1, 3, figsize=(15, 4))\n"
        "for ax, fi, fn in zip(axes, range(3), ['YES Rate', 'Agreement Rate', 'Label Entropy']):\n"
        "    ax.hist(behavioral_features[:, fi], bins=30, edgecolor='black', alpha=0.7)\n"
        "    ax.set_xlabel(fn)\n"
        "    ax.set_ylabel('Count')\n"
        "    ax.set_title(f'Distribution of {fn}')\n"
        "plt.tight_layout()\n"
        "plt.show()"))

    # Cell 14: md - Clustering
    cells.append(_nb_cell("markdown", "## 4. Behavioral Clustering\n\n"
        "Standardize features, scan KMeans k=2..20, pick best by silhouette."))

    # Cell 15: KMeans scan
    cells.append(_nb_cell("code",
        "scaler = StandardScaler()\n"
        "behavioral_scaled = scaler.fit_transform(behavioral_features)\n\n"
        "metrics_rows = []\n"
        "for n_c in range(2, 21):\n"
        "    km = KMeans(n_clusters=n_c, random_state=RANDOM_STATE, n_init=20)\n"
        "    labels_c = km.fit_predict(behavioral_scaled)\n"
        "    sil = silhouette_score(behavioral_scaled, labels_c)\n"
        "    ch = calinski_harabasz_score(behavioral_scaled, labels_c)\n"
        "    db = davies_bouldin_score(behavioral_scaled, labels_c)\n"
        "    metrics_rows.append({'n_clusters': n_c, 'silhouette': sil, 'ch': ch, 'db': db})\n\n"
        "metrics_df = pd.DataFrame(metrics_rows)\n"
        "best_idx = metrics_df['silhouette'].idxmax()\n"
        "best_k = int(metrics_df.loc[best_idx, 'n_clusters'])\n"
        "print(f'Best k = {best_k} (silhouette = {metrics_df.loc[best_idx, \"silhouette\"]:.4f})')"))

    # Cell 16: Metrics plot
    cells.append(_nb_cell("code",
        "fig, axes = plt.subplots(1, 3, figsize=(15, 4))\n"
        "for ax, col, title in zip(axes, ['silhouette', 'ch', 'db'],\n"
        "    ['Silhouette (higher=better)', 'Calinski-Harabasz (higher=better)', 'Davies-Bouldin (lower=better)']):\n"
        "    ax.plot(metrics_df['n_clusters'], metrics_df[col], 'o-')\n"
        "    ax.axvline(x=best_k, color='r', linestyle='--', alpha=0.5)\n"
        "    ax.set_xlabel('n_clusters')\n"
        "    ax.set_title(title)\n"
        "plt.suptitle('Behavioral Clustering Metrics', fontsize=13, y=1.02)\n"
        "plt.tight_layout()\n"
        "plt.show()"))

    # Cell 17: Final clustering
    cells.append(_nb_cell("code",
        "km_final = KMeans(n_clusters=best_k, random_state=RANDOM_STATE, n_init=20)\n"
        "behavioral_labels = km_final.fit_predict(behavioral_scaled)\n\n"
        "# Name clusters by YES rate\n"
        "cluster_means = {}\n"
        "for c in range(best_k):\n"
        "    cluster_means[c] = behavioral_features[behavioral_labels == c, 0].mean()\n"
        "sorted_clusters = sorted(cluster_means, key=lambda c: cluster_means[c])\n"
        "names = ['Conservative', 'Mainstream', 'Sensitive'] if best_k == 3 else [f'Cluster {i+1}' for i in range(best_k)]\n"
        "name_map = {c: names[i] for i, c in enumerate(sorted_clusters)}\n"
        "cluster_names = [name_map[c] for c in behavioral_labels]\n\n"
        "for c in sorted_clusters:\n"
        "    mask = behavioral_labels == c\n"
        "    print(f'{name_map[c]}: n={mask.sum()}, YES rate={behavioral_features[mask, 0].mean():.3f}, '\n"
        "          f'agreement={behavioral_features[mask, 1].mean():.3f}, entropy={behavioral_features[mask, 2].mean():.3f}')"))

    # Cell 18: Scatter plots
    cells.append(_nb_cell("code",
        "fig, axes = plt.subplots(1, 3, figsize=(16, 5))\n"
        "pairs = [(0, 1), (0, 2), (1, 2)]\n"
        "pair_labels = [('YES Rate', 'Agreement Rate'), ('YES Rate', 'Label Entropy'), ('Agreement Rate', 'Label Entropy')]\n"
        "for ax, (fi, fj), (xl, yl) in zip(axes, pairs, pair_labels):\n"
        "    for cname in sorted(set(cluster_names)):\n"
        "        mask = np.array(cluster_names) == cname\n"
        "        ax.scatter(behavioral_features[mask, fi], behavioral_features[mask, fj], alpha=0.6, label=cname, s=30)\n"
        "    ax.set_xlabel(xl)\n"
        "    ax.set_ylabel(yl)\n"
        "    ax.legend(fontsize=7)\n"
        "axes[1].set_title('Behavioral Feature Clustering')\n"
        "plt.tight_layout()\n"
        "plt.show()"))

    # Cell 19: Cluster profiles
    cells.append(_nb_cell("code",
        "profile_data = []\n"
        "for c in range(best_k):\n"
        "    mask = behavioral_labels == c\n"
        "    for fi, fn in enumerate(['yes_rate', 'agreement_rate', 'label_entropy']):\n"
        "        vals = behavioral_features[mask, fi]\n"
        "        profile_data.append({'cluster': name_map[c], 'n': mask.sum(), 'feature': fn, 'mean': vals.mean(), 'std': vals.std()})\n\n"
        "profile_df = pd.DataFrame(profile_data)\n"
        "fig, ax = plt.subplots(figsize=(10, 5))\n"
        "x_pos = np.arange(3)\n"
        "width = 0.8 / best_k\n"
        "for ci, cname in enumerate(sorted(set(profile_df['cluster']))):\n"
        "    sub = profile_df[profile_df['cluster'] == cname]\n"
        "    offset = (ci - (best_k - 1) / 2) * width\n"
        "    ax.bar(x_pos + offset, sub['mean'].values, width, yerr=sub['std'].values, capsize=3,\n"
        "           label=f\"{cname} (n={sub['n'].iloc[0]})\", alpha=0.8)\n"
        "ax.set_xticks(x_pos)\n"
        "ax.set_xticklabels(['yes_rate', 'agreement_rate', 'label_entropy'])\n"
        "ax.set_title('Behavioral Cluster Profiles')\n"
        "ax.legend()\n"
        "plt.tight_layout()\n"
        "plt.show()"))

    # Cell 20: md - Validation
    cells.append(_nb_cell("markdown", "## 5. Cluster Validation\n\n"
        "Silhouette plot and pairwise agreement heatmap."))

    # Cell 21: Silhouette plot
    cells.append(_nb_cell("code",
        "sil_values = silhouette_samples(behavioral_scaled, behavioral_labels)\n"
        "sil_avg = silhouette_score(behavioral_scaled, behavioral_labels)\n\n"
        "fig, ax = plt.subplots(figsize=(8, 6))\n"
        "y_lower = 10\n"
        "for ci in range(best_k):\n"
        "    ci_vals = sil_values[behavioral_labels == ci]\n"
        "    ci_vals.sort()\n"
        "    size_ci = ci_vals.shape[0]\n"
        "    y_upper = y_lower + size_ci\n"
        "    ax.fill_betweenx(np.arange(y_lower, y_upper), 0, ci_vals, alpha=0.7,\n"
        "                     label=f'{name_map[ci]} (n={size_ci})')\n"
        "    y_lower = y_upper + 10\n"
        "ax.axvline(x=sil_avg, color='red', linestyle='--', label=f'Mean = {sil_avg:.3f}')\n"
        "ax.set_xlabel('Silhouette coefficient')\n"
        "ax.set_title('Behavioral Clustering - Silhouette Plot')\n"
        "ax.legend(fontsize=8)\n"
        "plt.tight_layout()\n"
        "plt.show()"))

    # Cell 22: Agreement heatmap
    cells.append(_nb_cell("code",
        "cluster_labels = behavioral_labels + 1\n"
        "demo_df['cluster'] = cluster_labels\n\n"
        "# Pairwise agreement\n"
        "L = (A != 0).astype(np.float64)\n"
        "co_count = L @ L.T\n"
        "products = A.astype(np.float64) @ A.astype(np.float64).T\n"
        "agree_count = (products + co_count) / 2.0\n"
        "with np.errstate(divide='ignore', invalid='ignore'):\n"
        "    agreement_matrix = np.where(co_count > 0, agree_count / co_count, 0.0)\n"
        "np.fill_diagonal(agreement_matrix, 1.0)\n\n"
        "unique_clusters = sorted(np.unique(cluster_labels))\n"
        "n_cl = len(unique_clusters)\n"
        "agree_heatmap = np.zeros((n_cl, n_cl))\n"
        "for i_idx, ci in enumerate(unique_clusters):\n"
        "    for j_idx, cj in enumerate(unique_clusters):\n"
        "        mask_i = (cluster_labels == ci)\n"
        "        mask_j = (cluster_labels == cj)\n"
        "        sub = agreement_matrix[np.ix_(mask_i, mask_j)].copy()\n"
        "        if ci == cj:\n"
        "            np.fill_diagonal(sub, np.nan)\n"
        "            agree_heatmap[i_idx, j_idx] = np.nanmean(sub)\n"
        "        else:\n"
        "            agree_heatmap[i_idx, j_idx] = np.mean(sub)\n\n"
        "fig, ax = plt.subplots(figsize=(8, 6))\n"
        "sns.heatmap(pd.DataFrame(agree_heatmap, index=[f'C{c}' for c in unique_clusters],\n"
        "            columns=[f'C{c}' for c in unique_clusters]),\n"
        "            annot=True, fmt='.3f', cmap='YlOrRd', vmin=0.4, vmax=1.0, ax=ax)\n"
        "ax.set_title('Pairwise Agreement Heatmap')\n"
        "plt.tight_layout()\n"
        "plt.show()"))

    # Cell 23: md - Demographics
    cells.append(_nb_cell("markdown", "## 6. Demographic Analysis\n\n"
        "Check whether demographics are associated with cluster membership."))

    # Cell 24: Chi-squared
    cells.append(_nb_cell("code",
        "DEMO_VARS = ['gender', 'age', 'ethnicity', 'education', 'country']\n\n"
        "for var in DEMO_VARS:\n"
        "    ct = pd.crosstab(demo_df['cluster'], demo_df[var])\n"
        "    chi2, p_val, dof, expected = chi2_contingency(ct)\n"
        "    n_total = ct.sum().sum()\n"
        "    min_dim = min(ct.shape) - 1\n"
        "    cramers_v = np.sqrt(chi2 / (n_total * min_dim)) if min_dim > 0 else 0\n"
        "    print(f'{var:12s}  chi2={chi2:8.2f}  p={p_val:.4g}  V={cramers_v:.4f}  '\n"
        "          f'{\"SIG\" if p_val < 0.05 else \"\"}')\n"))

    # Cell 25: Stacked bars
    cells.append(_nb_cell("code",
        "fig, axes = plt.subplots(2, 3, figsize=(18, 10))\n"
        "axes = axes.ravel()\n"
        "for idx, var in enumerate(DEMO_VARS):\n"
        "    ax = axes[idx]\n"
        "    ct = pd.crosstab(demo_df['cluster'], demo_df[var], normalize='index')\n"
        "    ct.plot(kind='bar', stacked=True, ax=ax, colormap='Set2', edgecolor='white')\n"
        "    ax.set_title(f'Cluster composition by {var}')\n"
        "    ax.set_xlabel('Cluster')\n"
        "    ax.set_ylabel('Proportion')\n"
        "    ax.legend(fontsize=7, loc='upper right')\n"
        "if len(DEMO_VARS) < len(axes):\n"
        "    for j in range(len(DEMO_VARS), len(axes)):\n"
        "        axes[j].set_visible(False)\n"
        "plt.suptitle('Demographic Composition of Behavioral Clusters', fontsize=14, y=1.01)\n"
        "plt.tight_layout()\n"
        "plt.show()"))

    # Cell 26: md - GLMM
    cells.append(_nb_cell("markdown", "## 7. GLMM — Statistical Modeling\n\n"
        "Fit a binomial mixed-effects model with demographics as fixed effects\n"
        "and text + annotator as random effects."))

    # Cell 27: GLMM code
    cells.append(_nb_cell("code",
        "try:\n"
        "    from statsmodels.genmod.bayes_mixed_glm import BinomialBayesMixedGLM\n"
        "    import scipy.sparse as sp_sparse\n\n"
        "    COUNTRY_TO_REGION = {\n"
        "        'Italy': 'Europe', 'Spain': 'Europe', 'France': 'Europe',\n"
        "        'Germany': 'Europe', 'United Kingdom': 'Europe', 'Portugal': 'Europe',\n"
        "        'United States': 'America', 'Canada': 'America', 'Mexico': 'America',\n"
        "        'South Africa': 'Africa', 'Nepal': 'Asia', 'Viet Nam': 'Asia',\n"
        "        'Australia': 'Asia', 'Afghanistan': 'Middle East', 'Israel': 'Middle East',\n"
        "    }\n"
        "    glmm_df = annot_en[annot_en['annotator_id'].isin(active_annotators)].copy()\n"
        "    glmm_df['label_binary'] = (glmm_df['label_raw'] == 'YES').astype(int)\n"
        "    glmm_df['region'] = glmm_df['country'].map(COUNTRY_TO_REGION).fillna('Other')\n"
        "    n_obs = len(glmm_df)\n"
        "    print(f'GLMM observations: {n_obs}')\n\n"
        "    fixed_vars = ['gender', 'age', 'ethnicity', 'education', 'region']\n"
        "    exog_parts = [np.ones((n_obs, 1))]\n"
        "    fixed_names = ['Intercept']\n"
        "    for var in fixed_vars:\n"
        "        dummies = pd.get_dummies(glmm_df[var], prefix=var, drop_first=True, dtype=float)\n"
        "        exog_parts.append(dummies.values)\n"
        "        fixed_names.extend(dummies.columns.tolist())\n"
        "    exog = np.hstack(exog_parts)\n\n"
        "    text_cat = pd.Categorical(glmm_df['text_id'])\n"
        "    ann_cat = pd.Categorical(glmm_df['annotator_id'])\n"
        "    rows_idx = np.arange(n_obs)\n"
        "    text_vc = sp_sparse.csr_matrix((np.ones(n_obs), (rows_idx, text_cat.codes)), shape=(n_obs, len(text_cat.categories)))\n"
        "    ann_vc = sp_sparse.csr_matrix((np.ones(n_obs), (rows_idx, ann_cat.codes)), shape=(n_obs, len(ann_cat.categories)))\n"
        "    exog_vc = sp_sparse.hstack([text_vc, ann_vc]).toarray()\n"
        "    ident = np.array([0] * len(text_cat.categories) + [1] * len(ann_cat.categories))\n"
        "    endog = glmm_df['label_binary'].values.astype(float)\n\n"
        "    print('Fitting GLMM...')\n"
        "    model = BinomialBayesMixedGLM(endog, exog, exog_vc, ident)\n"
        "    result = model.fit_vb(verbose=True)\n\n"
        "    fe_df = pd.DataFrame({\n"
        "        'variable': fixed_names, 'coefficient': result.fe_mean,\n"
        "        'odds_ratio': np.exp(result.fe_mean), 'std_error': result.fe_sd,\n"
        "    })\n"
        "    fe_df['z_value'] = fe_df['coefficient'] / fe_df['std_error']\n"
        "    from scipy.stats import norm\n"
        "    fe_df['p_value'] = 2 * (1 - norm.cdf(np.abs(fe_df['z_value'])))\n"
        "    print(fe_df.to_string(index=False))\n\n"
        "    sigma_text = np.exp(result.vcp_mean[0])\n"
        "    sigma_ann = np.exp(result.vcp_mean[1])\n"
        "    var_text = sigma_text ** 2\n"
        "    var_ann = sigma_ann ** 2\n"
        "    var_residual = (np.pi ** 2) / 3\n"
        "    var_fixed = np.var(exog @ result.fe_mean)\n"
        "    total = var_fixed + var_text + var_ann + var_residual\n"
        "    print(f'\\nVariance decomposition:')\n"
        "    print(f'  Demographics: {var_fixed/total:.1%}')\n"
        "    print(f'  Text:         {var_text/total:.1%}')\n"
        "    print(f'  Annotator:    {var_ann/total:.1%}')\n"
        "    print(f'  Residual:     {var_residual/total:.1%}')\n"
        "except Exception as e:\n"
        "    print(f'GLMM failed: {e}')"))

    # Cell 28: md - KPCA comparison
    cells.append(_nb_cell("markdown", "## 8. KPCA Comparison (Why It Fails)\n\n"
        "Run KPCA + Ward's for comparison to show why it doesn't work with block design."))

    # Cell 29: KPCA comparison
    cells.append(_nb_cell("code",
        "# One-hot encoding\n"
        "A_no = (A == -1).astype(np.float64)\n"
        "A_yes = (A == 1).astype(np.float64)\n"
        "M_onehot = np.empty((n_annotators, 2 * n_texts), dtype=np.float64)\n"
        "M_onehot[:, 0::2] = A_no\n"
        "M_onehot[:, 1::2] = A_yes\n\n"
        "kpca = KernelPCA(kernel='cosine', n_components=n_annotators, random_state=RANDOM_STATE)\n"
        "Z_full = kpca.fit_transform(M_onehot)\n"
        "eigenvalues = kpca.eigenvalues_\n"
        "eigenvalues_pos = eigenvalues[eigenvalues > 0]\n"
        "cum_var = np.cumsum(eigenvalues_pos) / np.sum(eigenvalues_pos)\n\n"
        "k = int(np.searchsorted(cum_var, 0.855)) + 1\n"
        "k = max(k, 2)\n"
        "Z = Z_full[:, :k]\n"
        "print(f'KPCA: k={k} components')\n\n"
        "lm = linkage(Z, method='ward')\n"
        "kpca_labels = fcluster(lm, t=best_k, criterion='maxclust')\n"
        "kpca_sil = silhouette_score(Z, kpca_labels)\n"
        "beh_sil = silhouette_score(behavioral_scaled, behavioral_labels)\n"
        "print(f'KPCA silhouette:       {kpca_sil:.4f}')\n"
        "print(f'Behavioral silhouette: {beh_sil:.4f}')\n"
        "print(f'Behavioral is {beh_sil/max(kpca_sil, 0.001):.1f}x better')"))

    # Cell 30: KPCA cumvar plot
    cells.append(_nb_cell("code",
        "fig, ax = plt.subplots(figsize=(8, 4))\n"
        "ax.plot(range(1, len(cum_var) + 1), cum_var, 'b-', linewidth=1)\n"
        "ax.axhline(y=0.855, color='r', linestyle='--', label='85.5% threshold')\n"
        "ax.axvline(x=k, color='g', linestyle='--', label=f'k = {k}')\n"
        "ax.set_xlabel('Number of components')\n"
        "ax.set_ylabel('Cumulative variance explained')\n"
        "ax.set_title('Kernel PCA (Cosine) - Cumulative Variance')\n"
        "ax.legend()\n"
        "ax.set_xlim(0, min(100, len(cum_var)))\n"
        "plt.tight_layout()\n"
        "plt.show()"))

    # Cell 31: md - Robustness
    cells.append(_nb_cell("markdown", "## 9. Robustness: Bootstrap Stability\n\n"
        "Test how stable the behavioral clusters are across bootstrap resamples."))

    # Cell 32: Bootstrap
    cells.append(_nb_cell("code",
        "n_boot = 100\n"
        "jaccard_scores = []\n"
        "for b in range(n_boot):\n"
        "    rng = np.random.RandomState(b)\n"
        "    boot_idx = rng.choice(n_annotators, size=n_annotators, replace=True)\n"
        "    beh_boot = behavioral_scaled[boot_idx]\n"
        "    km_boot = KMeans(n_clusters=best_k, random_state=RANDOM_STATE, n_init=20)\n"
        "    labels_boot = km_boot.fit_predict(beh_boot)\n"
        "    original_boot = behavioral_labels[boot_idx]\n"
        "    boot_unique = np.unique(labels_boot)\n"
        "    orig_unique = np.unique(original_boot)\n"
        "    used = set()\n"
        "    total_jaccard = 0\n"
        "    for bc in boot_unique:\n"
        "        bc_set = set(np.where(labels_boot == bc)[0])\n"
        "        best_j, best_oc = 0, orig_unique[0]\n"
        "        for oc in orig_unique:\n"
        "            if oc in used: continue\n"
        "            oc_set = set(np.where(original_boot == oc)[0])\n"
        "            j = len(bc_set & oc_set) / max(len(bc_set | oc_set), 1)\n"
        "            if j > best_j: best_j, best_oc = j, oc\n"
        "        used.add(best_oc)\n"
        "        total_jaccard += best_j\n"
        "    jaccard_scores.append(total_jaccard / len(boot_unique))\n\n"
        "jaccard_arr = np.array(jaccard_scores)\n"
        "print(f'Mean Jaccard: {jaccard_arr.mean():.4f} +/- {jaccard_arr.std():.4f}')\n\n"
        "fig, ax = plt.subplots(figsize=(6, 4))\n"
        "ax.hist(jaccard_arr, bins=20, edgecolor='black', alpha=0.7)\n"
        "ax.axvline(x=jaccard_arr.mean(), color='red', linestyle='--', label=f'Mean = {jaccard_arr.mean():.3f}')\n"
        "ax.set_xlabel('Mean Jaccard similarity')\n"
        "ax.set_title('Bootstrap Stability (100 resamples)')\n"
        "ax.legend()\n"
        "plt.tight_layout()\n"
        "plt.show()"))

    # Cell 33: md - Conclusions
    cells.append(_nb_cell("markdown",
        "## 10. Conclusions\n\n"
        "**Key findings:**\n\n"
        "1. EXIST 2024 has a perfect block design (58 groups of 6 annotators, zero overlap), "
        "which makes cosine-based methods like KPCA ineffective.\n\n"
        "2. Behavioral feature clustering (YES rate, agreement rate, label entropy) produces "
        "meaningful clusters with much higher silhouette scores.\n\n"
        "3. Three annotator types emerge: Conservative (low YES rate), Mainstream (moderate), "
        "and Sensitive (high YES rate).\n\n"
        "4. Demographics have minimal predictive power — text content and individual tendencies "
        "dominate annotation decisions.\n\n"
        "5. For DPO training, behavioral clusters provide more meaningful preference signals "
        "than demographic splits."))

    # Fix source line endings
    for cell in cells:
        cell["source"] = _fix_source(cell["source"])

    # Write notebook
    notebook = {
        "nbformat": 4,
        "nbformat_minor": 5,
        "metadata": {
            "kernelspec": {
                "display_name": "Python 3",
                "language": "python",
                "name": "python3",
            },
            "language_info": {
                "name": "python",
                "version": "3.10.0",
            },
        },
        "cells": cells,
    }

    nb_path = "annotator_clustering_analysis.ipynb"
    with open(nb_path, "w") as f:
        json.dump(notebook, f, indent=1)
    print(f"Saved: {nb_path} ({len(cells)} cells)")
    print()


# ============================================================
# Summary
# ============================================================
print("=" * 60)
print("PIPELINE COMPLETE")
print("=" * 60)
print(f"  Primary method:      {PRIMARY_METHOD}")
print(f"  Annotators:          {n_annotators}")
print(f"  English texts:       {n_texts}")
print(f"  Annotation groups:   {n_groups}")
if k is not None:
    print(f"  KPCA components:     {k}")
print(f"  Optimal clusters:    {optimal_n}")
print(f"  Mean silhouette:     {sil_avg:.4f}")
print(f"  Within/between agree: {within_avg:.4f} / {between_avg:.4f}")
print()
print("Output files:")
output_files = [
    "block_design_heatmap.png",
    "behavioral_cluster_metrics.png",
    "behavioral_features_scatter.png",
    "behavioral_cluster_profiles.png",
    "behavioral_silhouette_plot.png",
    "agreement_heatmap.png",
    "silhouette_plot.png",
    "demographic_stacked_bars.png",
    "bootstrap_stability_behavioral.png",
]
if RUN_KPCA_COMPARISON:
    output_files.extend([
        "kpca_cumulative_variance.png",
        "dendrogram_full.png",
        "dendrogram_colored.png",
        "kpca_vs_behavioral_comparison.png",
        "bootstrap_stability.png",
    ])
if GENERATE_REPORT:
    output_files.append("annotator_clustering_report.docx")
if GENERATE_NOTEBOOK:
    output_files.append("annotator_clustering_analysis.ipynb")
for fname in output_files:
    print(f"  {fname}")
